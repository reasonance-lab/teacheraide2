# Description: Converts QTI XML questions into a Microsoft Word (.docx) document.
# file name: docx_converter.py

import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement # For OMML
from docx.oxml.ns import qn # For OMML namespaces
import io
import base64
from PIL import Image
import streamlit as st # For potential error reporting or progress
import re # Import regex for LaTeX detection

# Define namespace for QTI parsing
QTI_NS = {"qti": "http://www.imsglobal.org/xsd/imsqti_v2p2"}

# Simple regex to find potential LaTeX blocks and basic fractions
LATEX_PATTERN = re.compile(r"(\$.*?\$|\$\$[\s\S]*?\$\$|\\\(.*?\\\)|\\\[[\s\S]*?\\\])")
# Very basic fraction regex: \frac{num}{den} (doesn't handle nested braces well)
FRAC_PATTERN = re.compile(r"\\frac\{([^}]*)\}\{([^}]*)\}")

class QTIToDocxConverter:
    """Converts a list of QTI XML questions to a DOCX file."""

    def __init__(self, questions_xml: list, media_files: dict = None, title: str = "Quiz Paper"):
        """
        Initializes the converter.

        Args:
            questions_xml (list): A list of strings, each containing QTI XML for one question.
            media_files (dict): A dictionary where keys are media filenames (e.g., 'media/image1.png')
                                and values are the image data in bytes. Defaults to None.
            title (str): The title for the generated Word document.
        """
        self.questions_xml = questions_xml
        self.media_files = media_files if media_files else {}
        self.title = title
        self.doc = Document()
        # Add basic styling if needed (optional)
        # self.doc.styles['Normal'].font.name = 'Calibri'
        # self.doc.styles['Normal'].font.size = Pt(11)

    def _add_image(self, paragraph, image_filename):
        """Adds an image from media_files to the document, trying with and without 'media/' prefix."""
        image_key_to_try = image_filename
        image_data = None

        # Try the exact filename first
        if image_key_to_try in self.media_files:
            image_data = self.media_files[image_key_to_try]
        else:
            # If not found, try stripping 'media/' prefix
            if image_key_to_try.startswith("media/"):
                stripped_key = image_key_to_try[len("media/"):]
                if stripped_key in self.media_files:
                    image_key_to_try = stripped_key # Use the stripped key
                    image_data = self.media_files[image_key_to_try]

        if image_data:
            try:
                # image_data = self.media_files[image_filename] # Original line removed
                image_stream = io.BytesIO(image_data)
                # Optional: Resize image if needed
                # img = Image.open(image_stream)
                # width, height = img.size
                # max_width = Inches(4.0) # Example max width
                # if width > max_width * 96: # Assuming 96 DPI
                #     ratio = max_width / (width / 96)
                #     height = height * ratio
                #     width = max_width
                # else:
                #     width = Inches(width / 96)
                #     height = Inches(height / 96)
                # Reset stream position after reading with PIL
                image_stream.seek(0)
                paragraph.add_run().add_picture(image_stream, width=Inches(3.0)) # Adjust width as needed
            except Exception as e:
                st.warning(f"Could not add image '{image_filename}': {e}")
        else:
            # Add debugging output to show available keys when an image isn't found
            # Add debugging output only if image_data is still None after trying both keys
            st.warning(f"Image file '{image_filename}' not found in media_files (tried both with/without 'media/' prefix). Available keys: {list(self.media_files.keys())}")

    # --- OMML Generation Helpers ---
    def _create_omml_run(self, text):
        """Creates an OxmlElement for a simple text run within OMML."""
        r = OxmlElement('m:r')
        t = OxmlElement('m:t')
        t.text = text
        r.append(t)
        return r

    def _create_omml_fraction(self, num_text, den_text):
        """Creates an OxmlElement for a fraction."""
        f = OxmlElement('m:f')
        fPr = OxmlElement('m:fPr')
        type_val = OxmlElement('m:type')
        type_val.set(qn('m:val'), 'bar') # Standard fraction bar
        fPr.append(type_val)
        f.append(fPr)

        num = OxmlElement('m:num') # Numerator
        num.append(self._create_omml_run(num_text))
        f.append(num)

        den = OxmlElement('m:den') # Denominator
        den.append(self._create_omml_run(den_text))
        f.append(den)
        return f

    # --- End OMML Helpers ---

    def _parse_and_add_question(self, question_xml: str, question_number: int):
        """Parses a single QTI XML question and adds it to the DOCX document."""
        try:
            root = ET.fromstring(question_xml)
            identifier = root.get('identifier', f'Q{question_number}')
            q_title = root.get('title', f'Question {question_number}')

            # Find prompt
            prompt_elem = root.find(".//qti:prompt", QTI_NS)
            prompt_text = "".join(prompt_elem.itertext()).strip() if prompt_elem is not None else "No prompt found."

            # Add question number and prompt using the LaTeX helper
            p = self.doc.add_paragraph()
            p.add_run(f"{question_number}. ").bold = True
            self._add_text_with_latex(p, prompt_text) # Use helper here

            # Check for image(s) within the itemBody (more robust search)
            item_body = root.find(".//qti:itemBody", QTI_NS)
            if item_body is not None:
                # Find all images within itemBody, not just the first direct child
                img_elements = item_body.findall(".//qti:img", QTI_NS)
                for img_elem in img_elements:
                    img_src = img_elem.get("src")
                    if img_src:
                        # Add image below the prompt text
                        img_p = self.doc.add_paragraph()
                        self._add_image(img_p, img_src)
                        img_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else: # Fallback search if item_body structure is unusual
                 img_elem = root.find(".//qti:img", QTI_NS)
                 if img_elem is not None:
                    img_src = img_elem.get("src")
                    if img_src: # Indentation fixed here
                        # Add image below the prompt text
                        img_p = self.doc.add_paragraph()
                        self._add_image(img_p, img_src) # Corrected call placement
                        img_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # --- Improved logic for finding the main interaction ---
            interaction_elem = None
            possible_interactions = [
                "choiceInteraction", "orderInteraction", "textEntryInteraction",
                "extendedTextInteraction", "matchInteraction", "hotspotInteraction",
                "gapMatchInteraction", "inlineChoiceInteraction", "uploadInteraction"
                # Add other QTI interaction types if needed
            ]
            if item_body is not None:
                for tag_name in possible_interactions:
                    # Find the interaction element directly under itemBody or within a block element like div/p
                    interaction_elem = item_body.find(f".//qti:{tag_name}", QTI_NS)
                    if interaction_elem is not None:
                        break # Found the first main interaction

            interaction_type = interaction_elem.tag.split('}')[-1] if interaction_elem is not None else "unknown"

            # --- Route to specific handlers based on interaction type ---
            if interaction_type == "choiceInteraction":
                # Further check if it's True/False
                choices = interaction_elem.findall(".//qti:simpleChoice", QTI_NS)
                if len(choices) == 2 and all(c.get("identifier", "").lower() in ["true", "false", "t", "f"] for c in choices):
                     self._add_tf_question(root, interaction_elem)
                else:
                     self._add_choice_question(root, interaction_elem) # Pass interaction element
            elif interaction_type == "orderInteraction":
                self._add_order_question(root, interaction_elem)
            elif interaction_type == "textEntryInteraction":
                 # FIB might have multiple interactions within a <p> tag
                 fib_para_elem = item_body.find(".//qti:p[qti:textEntryInteraction]", QTI_NS) if item_body else None
                 if fib_para_elem is not None:
                     self._add_fib_question(root, fib_para_elem) # Pass the paragraph element
                 elif interaction_elem is not None: # Fallback if interaction is direct child
                     self._add_fib_question(root, interaction_elem)
                 else:
                     p = self.doc.add_paragraph(f"  (Could not find textEntryInteraction structure)")
                     p.italic = True
            elif interaction_type == "extendedTextInteraction":
                self._add_essay_question(root, interaction_elem)
            elif interaction_type == "matchInteraction":
                self._add_match_question(root, interaction_elem)
            # Add other types like hotspot, highlight, etc. as needed
            elif interaction_type != "unknown":
                 p = self.doc.add_paragraph(f"  (Question type '{interaction_type}' display not fully implemented)")
                 p.italic = True
            else:
                 # If no interaction found, maybe it's just text or unsupported structure
                 p = self.doc.add_paragraph(f"  (Could not determine question type or interaction element not found)")
                 p.italic = True

            self.doc.add_paragraph() # Add space after the question

        except ET.ParseError as e:
            st.error(f"Error parsing XML for question {question_number}: {e}")
            self.doc.add_paragraph(f"{question_number}. Error parsing question XML.")
        except Exception as e:
            st.error(f"Error processing question {question_number}: {e}")
            self.doc.add_paragraph(f"{question_number}. Error processing question.")

    def _add_choice_question(self, root, interaction_elem):
        """Adds MCQ or MRQ choices, handling text and images."""
        # Determine if multiple responses allowed (MRQ) vs single (MCQ)
        response_id = interaction_elem.get("responseIdentifier")
        response_decl = root.find(f".//qti:responseDeclaration[@identifier='{response_id}']", QTI_NS)
        is_multiple_response = response_decl is not None and response_decl.get("cardinality") == "multiple"

        choices = interaction_elem.findall(".//qti:simpleChoice", QTI_NS)
        for choice in choices:
            choice_id = choice.get("identifier")
            choice_text = "".join(choice.itertext()).strip() # Get all text content within the choice

            # Add choice identifier (e.g., A.)
            p = self.doc.add_paragraph(style='List Paragraph') # Use List Paragraph for better control
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25) # Hanging indent
            p.add_run(f"{choice_id}. ").bold = True

            # Add choice text using LaTeX helper
            self._add_text_with_latex(p, choice_text)

            # Check for image within the choice
            img_elem = choice.find(".//qti:img", QTI_NS)
            if img_elem is not None:
                img_src = img_elem.get("src")
                if img_src:
                    # Add image below the choice text
                    img_p = self.doc.add_paragraph()
                    self._add_image(img_p, img_src)
                    # Indent image paragraph slightly
                    img_p.paragraph_format.left_indent = Inches(0.5)

    def _add_order_question(self, root, interaction_elem):
        """Adds Order question choices (unshuffled), handling text."""
        choices = interaction_elem.findall(".//qti:simpleChoice", QTI_NS)
        self.doc.add_paragraph("  Arrange the following items in the correct order:")
        for i, choice in enumerate(choices):
            choice_id = choice.get("identifier")
            choice_text = "".join(choice.itertext()).strip()

            # Add choice identifier (e.g., A.)
            p = self.doc.add_paragraph(style='List Paragraph') # Use List Paragraph
            p.paragraph_format.left_indent = Inches(0.5) # Indent further than MCQ
            p.paragraph_format.first_line_indent = Inches(-0.25) # Hanging indent
            p.add_run(f"{chr(65+i)}. ").bold = False # Use A, B, C...

            # Add choice text using LaTeX helper
            self._add_text_with_latex(p, choice_text)
            # No images expected within order choices typically

    def _add_fib_question(self, root, fib_element):
        """Adds Fill-in-the-Blank question text with visual blanks, handling LaTeX.
           Accepts either the interaction element or the paragraph containing it.
        """
        para = self.doc.add_paragraph("  ") # Start with indentation

        # If the passed element is the interaction itself, find its parent paragraph
        if fib_element.tag.endswith("textEntryInteraction"):
             # Attempt to find the parent 'p' tag - this might be fragile
             item_body = root.find(".//qti:itemBody", QTI_NS)
             p_elem = item_body.find(".//qti:p[qti:textEntryInteraction]", QTI_NS) if item_body else None
             if p_elem is None: # Fallback if structure is different
                 p_elem = fib_element # Process just the interaction? Less ideal.
        else: # Assume it's the paragraph element
             p_elem = fib_element

        if p_elem is None:
             para.add_run("[Fill-in-the-blank text structure not found]")
             return

        # Process the paragraph content node by node
        if p_elem.text:
            self._add_text_with_latex(para, p_elem.text)

        for child in p_elem:
            if child.tag.endswith('textEntryInteraction'):
                # Add a visual blank (e.g., underscores)
                blank_length = int(child.get("expectedLength", 15))
                para.add_run(" " + "_" * blank_length + " ")
            else:
                # Add text from other elements if needed (like <span>, <a> etc.)
                 if child.text:
                     # Recursively handle potential LaTeX within other tags? For now, simple text.
                     # self._add_text_with_latex(para, child.text)
                     para.add_run(child.text) # Simpler for now

            # Add tail text after the child element
            if child.tail:
                self._add_text_with_latex(para, child.tail)


    def _add_essay_question(self, root, interaction_elem):
        """Adds Essay question prompt and space for answer."""
        expected_lines = int(interaction_elem.get("expectedLines", 5))
        # Add lines for the student to write on
        self.doc.add_paragraph("  Answer:")
        for _ in range(expected_lines):
            self.doc.add_paragraph("  " + "_" * 60) # Add visual lines

    def _add_match_question(self, root, interaction_elem):
        """Adds Matching question columns, handling text."""
        match_sets = interaction_elem.findall(".//qti:simpleMatchSet", QTI_NS)
        if len(match_sets) >= 2:
            source_choices_elem = match_sets[0].findall(".//qti:simpleAssociableChoice", QTI_NS)
            target_choices_elem = match_sets[1].findall(".//qti:simpleAssociableChoice", QTI_NS)

            # Extract text using helper function to handle potential LaTeX
            source_choices = {}
            for choice in source_choices_elem:
                 source_choices[choice.get("identifier")] = "".join(choice.itertext()).strip()

            target_choices = {}
            for choice in target_choices_elem:
                 target_choices[choice.get("identifier")] = "".join(choice.itertext()).strip()

            self.doc.add_paragraph("  Match the items in Column A with the items in Column B:")

            # Basic table approach (could be improved with actual table object)
            max_len = max(len(source_choices), len(target_choices))
            source_keys = list(source_choices.keys())
            target_keys = list(target_choices.keys())

            # Use docx table for better alignment
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid' # Optional: Add borders
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Column A'
            hdr_cells[1].text = 'Column B'
            # Set header bold (optional)
            for cell in hdr_cells:
                 cell.paragraphs[0].runs[0].bold = True

            # Add choices to table
            max_len = max(len(source_choices), len(target_choices))
            source_keys = list(source_choices.keys())
            target_keys = list(target_choices.keys())

            for i in range(max_len):
                row_cells = table.add_row().cells
                # Add source choice with LaTeX handling
                if i < len(source_keys):
                     p_source = row_cells[0].paragraphs[0] # Get existing paragraph
                     p_source.add_run(f"{i+1}. ")
                     self._add_text_with_latex(p_source, source_choices[source_keys[i]])
                # Add target choice with LaTeX handling
                if i < len(target_keys):
                     p_target = row_cells[1].paragraphs[0] # Get existing paragraph
                     p_target.add_run(f"{chr(65+i)}. ")
                     self._add_text_with_latex(p_target, target_choices[target_keys[i]])

            # Adjust column widths (optional)
            # table.columns[0].width = Inches(2.5)
            # table.columns[1].width = Inches(2.5)

        else:
             self.doc.add_paragraph("  [Matching sets not found or incomplete]")

    # --- Add the new True/False handler ---
    def _add_tf_question(self, root, interaction_elem):
         """Adds True/False choices."""
         # Identifiers are usually 'true'/'false' or 'T'/'F'
         p_true = self.doc.add_paragraph("  True", style='List Bullet')
         p_false = self.doc.add_paragraph("  False", style='List Bullet')

    # --- Add the LaTeX helper method (Display original LaTeX block as italic) ---
    def _add_text_with_latex(self, paragraph, text):
        """Adds text to a paragraph, formatting the original detected LaTeX block as italic."""
        if not text:
            return
        parts = LATEX_PATTERN.split(text)
        for i, part in enumerate(parts):
            if not part:
                continue

            is_latex = (i % 2 == 1) # Check if it's an odd-indexed part (LaTeX)

            # Add the original part (including delimiters if it's LaTeX)
            run = paragraph.add_run(part)

            # Apply italics if it was identified as LaTeX
            if is_latex:
                run.italic = True


    def generate_docx_bytes(self) -> bytes:
        """Generates the DOCX file content as bytes."""
        # Add title
        self.doc.add_heading(self.title, level=1)
        self.doc.add_paragraph() # Add space after title

        # Add instructions (optional)
        # self.doc.add_paragraph("Instructions: Please answer all questions to the best of your ability.")
        # self.doc.add_paragraph()

        # Process each question
        for i, q_xml in enumerate(self.questions_xml):
            self._parse_and_add_question(q_xml, i + 1)

        # Save to a byte stream
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

# --- Example Usage (for testing or integration) ---
# if __name__ == '__main__':
#     # This part would be replaced by integration into your Streamlit app
#     # Example QTI XML strings (replace with actual data)
#     sample_mcq = """<?xml version="1.0" encoding="UTF-8"?>
#     <assessmentItem xmlns="http://www.imsglobal.org/xsd/imsqti_v2p2" identifier="mcq1" title="Sample MCQ">
#         <responseDeclaration identifier="RESPONSE" cardinality="single" baseType="identifier">
#             <correctResponse><value>B</value></correctResponse>
#         </responseDeclaration>
#         <itemBody>
#             <choiceInteraction responseIdentifier="RESPONSE" shuffle="false" maxChoices="1">
#                 <prompt>What is the capital of France?</prompt>
#                 <simpleChoice identifier="A">London</simpleChoice>
#                 <simpleChoice identifier="B">Paris</simpleChoice>
#                 <simpleChoice identifier="C">Berlin</simpleChoice>
#             </choiceInteraction>
#         </itemBody>
#     </assessmentItem>"""

#     sample_fib = """<?xml version="1.0" encoding="UTF-8"?>
#     <assessmentItem xmlns="http://www.imsglobal.org/xsd/imsqti_v2p2" identifier="fib1" title="Sample FIB">
#         <responseDeclaration identifier="RESPONSE1" cardinality="single" baseType="string">
#             <correctResponse><value>red</value></correctResponse>
#         </responseDeclaration>
#          <responseDeclaration identifier="RESPONSE2" cardinality="single" baseType="string">
#             <correctResponse><value>blue</value></correctResponse>
#         </responseDeclaration>
#         <itemBody>
#             <p>Roses are <textEntryInteraction responseIdentifier="RESPONSE1" expectedLength="10"/> and violets are <textEntryInteraction responseIdentifier="RESPONSE2" expectedLength="10"/>.</p>
#         </itemBody>
#     </assessmentItem>"""

#     questions = [sample_mcq, sample_fib]
#     converter = QTIToDocxConverter(questions, title="My Sample Quiz")
#     docx_bytes = converter.generate_docx_bytes()

#     # Save the file locally for testing
#     with open("work_desk/output/sample_quiz.docx", "wb") as f:
#         f.write(docx_bytes)
#     print("Sample quiz saved to work_desk/output/sample_quiz.docx")
