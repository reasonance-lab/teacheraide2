import base64
from io import BytesIO
import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
import tempfile

from prompts.qti_prompts import create_hlt_prompt, create_lessonprep_prompt
from utils.llm_handlers import (
    generate_hlt_lesson_prep_anthropic_streaming,
    generate_hlt_lesson_prep_openai_streaming
)




def apply_text_formatting(text: str, patterns: dict) -> str:
    """Apply bold and italic formatting to text"""
    # Replace bold text
    text = patterns['bold'].sub(r"\1", text)
    
    # Replace italic text
    text = patterns['italic'].sub(r"\1", text)
    
    return text

def create_styles(doc):
    """Create custom styles for the document"""

    # Course title style
    course_style = doc.styles.add_style('Course', WD_STYLE_TYPE.PARAGRAPH)
    course_style.font.name = 'Calibri'
    course_style.font.size = Pt(11)
    course_style.font.italic = True
    course_style.font.color.rgb = RGBColor(65, 105, 225)  # Royal Blue

    # Section header style
    section_style = doc.styles.add_style('Section', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.name = 'Calibri'
    section_style.font.size = Pt(12)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor(65, 105, 225)  # Royal Blue

    # Subsection style
    subsection_style = doc.styles.add_style('Subsection', WD_STYLE_TYPE.PARAGRAPH)
    subsection_style.font.name = 'Calibri'
    subsection_style.font.size = Pt(11)
    subsection_style.font.bold = True

    # Note style
    note_style = doc.styles.add_style('Note', WD_STYLE_TYPE.PARAGRAPH)
    note_style.font.name = 'Calibri'
    note_style.font.size = Pt(11)
    note_style.font.italic = True
    note_style.font.color.rgb = RGBColor(128, 128, 128)  # Gray

# def markdown_to_word(markdown_content: str, file_name: str) -> str:
#     """
#     Convert markdown content to a formatted Word document.
    
#     Args:
#         markdown_content (str): The input markdown content.
#         file_name (str): The name of the output Word file.
    
#     Returns:
#         str: The file path of the created Word document.
#     """
#     import re
#     doc = Document()
#     create_styles(doc)  # Create custom styles

#     # Define regex patterns
#     patterns = {
#         # 'title': re.compile(r"^#\s+([^#]+)"),  # Single # for title
#         # 'course': re.compile(r"^\*\*Course\*\*:\s*(.*)"),
#         # 'section': re.compile(r"^##\s+([^#]+)"),  # Changed to ## for sections
#         # 'subsubsection': re.compile(r"^###\s+([^#]+)"), # <-- ADDED pattern for ###
#         # 'subsection': re.compile(r"^\*\*([^*]+)\*\*:\s*(.*)"),       #,re.compile(r"^###?\s+([^#]+)")
#         # 'bullet': re.compile(r"^(\s*)[•-]\s+(.*)"),
#         # 'numbered': re.compile(r"^(\s*)\d+\.\s+(.*)"),
#         # 'note': re.compile(r"^\*\*([^:]+):\*\*\s*(.*)"),
#         # 'bold': re.compile(r"\*\*([^*]+)\*\*"),
#         # 'italic': re.compile(r"\*([^*]+)\*")
#         'title': re.compile(r"^#\s+([^#]+)"),  # Single # for title
#         'course': re.compile(r"^\*\*Course\*\*:\s*(.*)"),
#         'section': re.compile(r"^##\s+([^#]+)"),  # ## for sections
#         'subsubsection': re.compile(r"^###\s+([^#]+)"), # <-- ADDED pattern for ###
#         'subsection': re.compile(r"^\*\*([^*]+)\*\*:\s*(.*)"), # For **Header**: style
#         'bullet': re.compile(r"^(\s*)[•*-]\s+(.*)"), # Adjusted to include '*' as bullet marker
#         'numbered': re.compile(r"^(\s*)\d+\.\s+(.*)"),
#         'note': re.compile(r"^\*\*([^:]+):\*\*\s*(.*)"), # Matches **Note:**, **Example:** etc.
#         'bold': re.compile(r"\*\*([^*]+)\*\*"),
#         'italic': re.compile(r"\*([^*]+)\*")
#     }

#     # Track list state
#     current_list_level = 0
#     in_list = False

#     lines = markdown_content.splitlines()
#     i = 0
#     while i < len(lines):
#         line = lines[i].rstrip()
        
#         # Skip empty lines
#         if not line:
#             i += 1
#             continue
        
#         # Title (single #)
#         if match := patterns['title'].match(line):
#             p = doc.add_paragraph(style='Title')
#             p.add_run(match.group(1).strip())
#             i += 1
#             continue

#         # Course information
#         if match := patterns['course'].match(line):
#             p = doc.add_paragraph(style='Course')
#             p.add_run(match.group(1))
#             i += 1
#             continue

#         # Section headers
#         if match := patterns['section'].match(line):
#             p = doc.add_paragraph(style='Section')
#             p.add_run(match.group(1).strip())
#             i += 1
#             continue

#         # Subsection headers
#         if match := patterns['subsection'].match(line):
#             p = doc.add_paragraph(style='Subsection')
#             p.add_run(f"{match.group(1)}: {match.group(2)}")
#             i += 1
#             continue

#         # Bullet points and numbered lists
#         bullet_match = patterns['bullet'].match(line)
#         number_match = patterns['numbered'].match(line)
        
#         if bullet_match or number_match:
#             match = bullet_match or number_match
#             indent_level = len(match.group(1)) // 2  # Calculate indent level
#             content = match.group(2)

#             # Choose style based on list type and level
#             style = 'List Number' if number_match else 'List Bullet'
#             if indent_level > 0:
#                 style += f' {indent_level + 1}'  # Word uses 1-based levels

#             # Add paragraph with appropriate style
#             p = doc.add_paragraph(style=style)
            
#             # Apply formatting to content
#             content = apply_text_formatting(content, patterns)
#             p.add_run(content)
            
#             i += 1
#             continue

#         # Notes and special sections
#         if match := patterns['note'].match(line):
#             p = doc.add_paragraph(style='Note')
#             p.add_run(f"{match.group(1)}: {match.group(2)}")
#             i += 1
#             continue

#         # Regular paragraph
#         p = doc.add_paragraph()
#         formatted_text = apply_text_formatting(line, patterns)
#         p.add_run(formatted_text)
#         i += 1

#     # Save document
#     temp_dir = tempfile.gettempdir()
#     file_path = f"{temp_dir}/{file_name}"
#     doc.save(file_path)
#     return file_path

def markdown_to_word(markdown_content: str, file_name: str) -> str:
    """
    Convert markdown content to a formatted Word document.

    Args:
        markdown_content (str): The input markdown content.
        file_name (str): The name of the output Word file.

    Returns:
        str: The file path of the created Word document.
    """
    import re
    doc = Document()
    create_styles(doc)  # Create custom styles

    # Define regex patterns
    patterns = {
        'title': re.compile(r"^#\s+([^#]+)"),  # Single # for title
        'course': re.compile(r"^\*\*Course\*\*:\s*(.*)"),
        'section': re.compile(r"^##\s+([^#]+)"),  # ## for sections
        'subsubsection': re.compile(r"^###\s+([^#]+)"), # <-- ADDED pattern for ###
        'subsection': re.compile(r"^\*\*([^*]+)\*\*:\s*(.*)"), # For **Header**: style
        'bullet': re.compile(r"^(\s*)[•*-]\s+(.*)"), # Adjusted to include '*' as bullet marker
        'numbered': re.compile(r"^(\s*)\d+\.\s+(.*)"),
        'note': re.compile(r"^\*\*([^:]+):\*\*\s*(.*)"), # Matches **Note:**, **Example:** etc.
        'bold': re.compile(r"\*\*([^*]+)\*\*"),
        'italic': re.compile(r"\*([^*]+)\*")
    }

    # Track list state (Although not used in the simplified logic below, kept for context)
    # current_list_level = 0
    # in_list = False # Removed as list handling logic simplified

    lines = markdown_content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Title (single #)
        if match := patterns['title'].match(line):
            p = doc.add_paragraph(style='Title')
            # Apply nested formatting if any (e.g., # Title with *italic*)
            formatted_title = apply_text_formatting(match.group(1).strip(), patterns)
            p.add_run(formatted_title)
            i += 1
            continue

        # Course information
        if match := patterns['course'].match(line):
            p = doc.add_paragraph(style='Course')
            formatted_text = apply_text_formatting(match.group(1).strip(), patterns)
            p.add_run(formatted_text)
            i += 1
            continue

        # Section headers (##)
        if match := patterns['section'].match(line):
            p = doc.add_paragraph(style='Section')
            formatted_text = apply_text_formatting(match.group(1).strip(), patterns)
            p.add_run(formatted_text)
            i += 1
            continue

        # Sub-section headers (###) <-- ADDED BLOCK
        if match := patterns['subsubsection'].match(line):
            p = doc.add_paragraph(style='Subsection') # Apply Subsection style
            formatted_text = apply_text_formatting(match.group(1).strip(), patterns)
            p.add_run(formatted_text)
            i += 1
            continue

        # Subsection headers (**Bold:**) - Kept this distinct, might be different intent
        if match := patterns['subsection'].match(line):
            p = doc.add_paragraph(style='Subsection') # Also apply Subsection style
            # Combine bolded part and the rest, then format potential inner markup
            combined_text = f"{match.group(1)}: {match.group(2)}"
            formatted_text = apply_text_formatting(combined_text, patterns)
            # Re-apply bold to the header part specifically if needed, but style handles it
            # For simplicity, let the style handle the bolding of the whole line
            p.add_run(formatted_text) # Style applies bold
            i += 1
            continue

        # Bullet points and numbered lists
        bullet_match = patterns['bullet'].match(line)
        number_match = patterns['numbered'].match(line)

        if bullet_match or number_match:
            match = bullet_match or number_match
            indent_level = len(match.group(1)) // 2 if match.group(1) else 0 # Handle no indent case
            content = match.group(2)

            # Choose style based on list type and level
            style_name = 'List Number' if number_match else 'List Bullet'
            # Word's default list styles handle indentation levels automatically up to a point
            # Explicitly setting List Bullet 2, 3 etc. might require ensuring those styles exist
            # Using the base style often works well for simple nesting
            if indent_level > 0:
                 # Try using numbered styles like 'List Number 2', 'List Number 3' etc.
                 # Requires these styles to be defined or exist in the default template
                 style_name += f' {indent_level + 1}'
                 # Fallback or safer approach: Use base style and rely on indentation property
                 # style_name = 'List Number' if number_match else 'List Bullet' # Use base style

            # Add paragraph with appropriate style
            # Check if style exists, fallback to default if not needed for this specific case
            try:
                 p = doc.add_paragraph(style=style_name)
            except KeyError: # If style like 'List Bullet 2' doesn't exist
                 print(f"Warning: Style '{style_name}' not found. Using base list style.")
                 base_style_name = 'List Number' if number_match else 'List Bullet'
                 p = doc.add_paragraph(style=base_style_name)


            # Set indentation (alternative/complementary to using numbered styles)
            # This provides visual indentation even if specific level styles aren't used/defined
            if indent_level > 0:
                 p.paragraph_format.left_indent = Pt(18 * indent_level) # Adjust multiplier as needed

            # Apply formatting to content
            content = apply_text_formatting(content, patterns)
            p.add_run(content)

            i += 1
            continue

        # Notes and special sections (**Note:**, **Reminder:** etc.)
        # Adjusted regex slightly to capture the label correctly
        note_match = re.match(r"^\*\*([A-Za-z\s]+):\*\*\s*(.*)", line) # More specific label match
        if note_match:
             p = doc.add_paragraph(style='Note')
             label = note_match.group(1)
             text = note_match.group(2)
             # Apply formatting to the text part
             formatted_text = apply_text_formatting(text, patterns)
             # Add label in bold (as per markdown) and text with formatting
             p.add_run(f"{label}: ").bold = True # Style is italic gray, make label bold explicitly
             p.add_run(formatted_text) # Text part inherits Note style (italic gray)
             # Or just let the Note style apply fully:
             # combined_note = f"{label}: {text}"
             # formatted_note = apply_text_formatting(combined_note, patterns)
             # p.add_run(formatted_note) # This might make the label italic+gray too depending on apply_text_formatting

             i += 1
             continue


        # Regular paragraph
        p = doc.add_paragraph()
        formatted_text = apply_text_formatting(line, patterns)
        p.add_run(formatted_text)
        i += 1

    # Save document
    temp_dir = tempfile.gettempdir()
    # Ensure file_name ends with .docx
    if not file_name.lower().endswith('.docx'):
        file_name += '.docx'
    file_path = f"{temp_dir}/{file_name}"
    doc.save(file_path)
    return file_path



def process_pdf_for_Claude(pdf_output: BytesIO) -> str:
    """Process PDF for Claude API, with size checks and compression if needed"""
    try:
        # Reset buffer position
        pdf_output.seek(0)
        # Get the PDF content
        pdf_content = pdf_output.getvalue()
        # Check original size
        original_size_mb = len(pdf_content) / (1024 * 1024)
        if original_size_mb > 10:  # If larger than 10MB
            # st.warning(f"PDF size ({original_size_mb:.2f}MB) is large, compressing...")
            try:
                import fitz  # PyMuPDF
                # Create PDF document from bytes
                doc = fitz.open(stream=pdf_content, filetype="pdf")
                # Create new PDF with compression
                output = BytesIO()
                doc.save(output, 
                        garbage=4,     # Max garbage collection
                        deflate=True,  # Use deflate compression
                        ascii=False,   # Allow binary content
                        linear=False)  # Non-linear PDF to save space
                
                # Get compressed content
                pdf_content = output.getvalue()
                compressed_size_mb = len(pdf_content) / (1024 * 1024)
                # st.write(f"Compressed PDF size: {compressed_size_mb:.2f}MB")
                doc.close()             
            except Exception as e:
                st.warning(f"Compression failed: {str(e)}. Using original PDF.")
        
        # Convert to base64
        base_64_encoded_data = base64.b64encode(pdf_content)
        pdf_data = base_64_encoded_data.decode('utf-8')
        final_size_mb = len(pdf_data.encode('utf-8')) / (1024 * 1024)
        # st.write(f"Final PDF data size: {final_size_mb:.2f}MB")
        return pdf_data
        
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
        return ""



def lesson_prep():
    """
    Lesson Preparation Tool for generating lesson plans based on High Leverage Tasks (HLTs).
    """
    # Check if the user has uploaded a PDF and processed it
    if 'extracted_pdf' not in st.session_state or st.session_state.extracted_pdf is None:
        st.warning("⚠️ Please upload and process a PDF first")
        return

    # Set the title of the page
    st.subheader(":material/fact_check: Lesson Preparation Tool", divider=True)
    
    if 'extracted_pdf' not in st.session_state:
        st.warning("⚠️ Please upload and process a PDF first")
        return
    
    # Initialize session state variables
    st.session_state.setdefault("hlt_response", "")
    st.session_state.setdefault("modified_hlt", "")
    st.session_state.setdefault("lessonprep_response", "")
    st.session_state.setdefault("modified_lessonprep", "")
    st.session_state.setdefault("hlt_confirmed", False)
    
    # Divide the page into two columns
    col1, col2 = st.columns(2, border=True)
    
    # Display HLT(s) in the left column
    with col1:
        st.subheader("High Leverage Tasks (HLTs)", divider=True)
        # Create the HLT text area that persists
        if st.session_state.hlt_response:
            st.session_state.modified_hlt = st.text_area(
                "Edit HLT(s) below:",
                value=st.session_state.modified_hlt or st.session_state.hlt_response,
                height=300,
                key="hlt_text_area"
            )
        
        if st.button("Generate High Leverage Task(s)"):
            with st.spinner("Generating High Leverage Task(s)"):
                if st.session_state.model_choice == "GPT-4o (OpenAI)" and st.session_state.user_openAIapi_key:
                    #encoded_images = process_pdf_for_openai(st.session_state['extracted_pdf'])
                    #st.session_state.encoded_images = encoded_images
                    pdf_content = process_pdf_for_Claude(st.session_state['extracted_pdf'])
                    
                    if pdf_content:
                        hlt_response = generate_hlt_lesson_prep_openai_streaming(
                            prompt=create_hlt_prompt("Honors Chemistry, high school"),
                            aimed_output="High Leverage Task(s)",
                            source_content="PDF content",
                            api_key=st.session_state.user_openAIapi_key,
                            pdf_content=pdf_content
                        )
                        if hlt_response:
                            st.session_state.hlt_response = hlt_response
                            st.session_state.modified_hlt = hlt_response
                            st.rerun()
                    else:
                        st.error("Failed to process PDF for OpenAI")
                        return

                elif st.session_state.model_choice == "Claude 3.7 Sonnet (Anthropic)" and st.session_state.user_anthropic_key:
                    pdf_content = process_pdf_for_Claude(st.session_state['extracted_pdf'])
                    st.session_state.pdf_content=pdf_content
                    hlt_response = generate_hlt_lesson_prep_anthropic_streaming(
                        create_hlt_prompt("Honors Chemistry, high school"),
                        pdf_content=pdf_content,
                        api_key=st.session_state.user_anthropic_key,
                        aimed_output="High Leverage Task(s)",
                        source_content="PDF content",
                        model="claude-3-5-sonnet-20241022"
                    )
                    if hlt_response:
                        st.session_state.hlt_response = hlt_response
                        st.session_state.modified_hlt = hlt_response
                        st.rerun()
                else:
                    st.error("Please configure API keys in the Settings tab")
                    return

    with col2:
        st.subheader("Lesson Preparation", divider=True)
        
        # Create persistent text area for lesson prep
        if st.session_state.lessonprep_response:
            st.session_state.modified_lessonprep = st.text_area(
                "Edit Lesson Preparation below:",
                value=st.session_state.modified_lessonprep or st.session_state.lessonprep_response,
                height=300,
                key="lessonprep_text_area"
            )
        
        # Confirm HLT button
        if st.button("Confirm HLT and generate lesson prep"):
            with st.spinner("Generating Lesson Preparation..."):
                if st.session_state.model_choice == "GPT-4o (OpenAI)":
                    # encoded_images = process_pdf_for_openai(st.session_state['extracted_pdf'])
                    # st.session_state.encoded_images = encoded_images
                    pdf_content = process_pdf_for_Claude(st.session_state['extracted_pdf'])
                    if pdf_content:
                        modified_hlt = st.session_state.modified_hlt
                        lessonprep_response = generate_hlt_lesson_prep_openai_streaming(
                            prompt=create_lessonprep_prompt(
                                subject_name="Honors Chemistry, high school", 
                                hlt=modified_hlt
                            ),
                            aimed_output="Lesson Preparation",
                            source_content="High Leverage Task(s) and then PDF content",
                            api_key=st.session_state.user_openAIapi_key,
                            pdf_content=pdf_content
                        )
                        if lessonprep_response:
                            st.session_state.lessonprep_response = lessonprep_response
                            st.session_state.modified_lessonprep = lessonprep_response
                            st.rerun()

                elif st.session_state.model_choice == "Claude 3.7 Sonnet (Anthropic)":
                    modified_hlt = st.session_state.modified_hlt
                    lessonprep_response = generate_hlt_lesson_prep_anthropic_streaming(
                        create_lessonprep_prompt("Honors Chemistry, high school", 
                                hlt=modified_hlt),
                        pdf_content=st.session_state.pdf_content,
                        api_key=st.session_state.user_anthropic_key,
                        aimed_output="Lesson Preparation",
                        source_content="High Leverage Task(s) and then PDF content",
                        model="claude-3-5-sonnet-20241022"
                    )
                    if lessonprep_response:
                        st.session_state.lessonprep_response = lessonprep_response
                        st.session_state.modified_lessonprep = lessonprep_response
                        st.rerun()
                else:
                    st.error("Please configure API keys in the Settings tab")
                    return

        # Download button
        if st.button("Download Lesson Prep"):
            try:
                # Use the modified lesson prep content
                file_path = markdown_to_word(
                    st.session_state.modified_lessonprep, 
                    "lesson_prep.docx"
                )
                with open(file_path, "rb") as file:
                    st.download_button(
                        label="Download as Word Document",
                        data=file,
                        file_name="lesson_prep.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"Error generating Word document: {e}")


lesson_prep()